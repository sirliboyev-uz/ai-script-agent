"""Export service for scripts."""
import io
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer


class ExportService:
    """Service for exporting scripts to various formats."""

    @staticmethod
    def export_to_txt(script_data: Dict[str, Any]) -> bytes:
        """
        Export script to plain text format.

        Args:
            script_data: Script data dictionary

        Returns:
            Text content as bytes
        """
        lines = []

        # Header
        lines.append("=" * 80)
        lines.append(script_data.get("title", "YouTube Script").upper())
        lines.append("=" * 80)
        lines.append("")

        # Metadata
        lines.append(f"Topic: {script_data.get('topic', 'N/A')}")
        lines.append(f"Style: {script_data.get('tone', 'N/A')}")
        lines.append(f"Duration: {script_data.get('estimated_duration', 'N/A')}")
        lines.append("")

        # Description
        if script_data.get("description"):
            lines.append("DESCRIPTION")
            lines.append("-" * 80)
            lines.append(script_data.get("description"))
            lines.append("")

        # Keywords
        if script_data.get("keywords"):
            lines.append("KEYWORDS")
            lines.append("-" * 80)
            lines.append(", ".join(script_data.get("keywords", [])))
            lines.append("")

        # Full Script
        lines.append("SCRIPT")
        lines.append("=" * 80)
        lines.append("")
        lines.append(script_data.get("full_script", ""))

        content = "\n".join(lines)
        return content.encode('utf-8')

    @staticmethod
    def export_to_docx(script_data: Dict[str, Any]) -> bytes:
        """
        Export script to Word document format.

        Args:
            script_data: Script data dictionary

        Returns:
            DOCX content as bytes
        """
        doc = Document()

        # Title
        title = doc.add_heading(script_data.get("title", "YouTube Script"), 0)
        title.alignment = 1  # Center

        # Metadata table
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Light Grid Accent 1'

        table.cell(0, 0).text = "Topic"
        table.cell(0, 1).text = script_data.get("topic", "N/A")

        table.cell(1, 0).text = "Style"
        table.cell(1, 1).text = script_data.get("tone", "N/A")

        table.cell(2, 0).text = "Duration"
        table.cell(2, 1).text = script_data.get("estimated_duration", "N/A")

        doc.add_paragraph()

        # Description
        if script_data.get("description"):
            doc.add_heading("Description", 1)
            doc.add_paragraph(script_data.get("description"))

        # Keywords
        if script_data.get("keywords"):
            doc.add_heading("Keywords", 1)
            doc.add_paragraph(", ".join(script_data.get("keywords", [])))

        # Script
        doc.add_heading("Script", 1)
        script_text = script_data.get("full_script", "")

        # Split into paragraphs for better formatting
        for paragraph in script_text.split("\n\n"):
            if paragraph.strip():
                p = doc.add_paragraph(paragraph.strip())
                p.style.font.size = Pt(12)

        # Save to bytes
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return file_stream.read()

    @staticmethod
    def export_to_pdf(script_data: Dict[str, Any]) -> bytes:
        """
        Export script to PDF format.

        Args:
            script_data: Script data dictionary

        Returns:
            PDF content as bytes
        """
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter,
                                rightMargin=72, leftMargin=72,
                                topMargin=72, bottomMargin=18)

        # Container for 'Flowable' objects
        elements = []

        # Define styles
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='CustomTitle',
                                  parent=styles['Heading1'],
                                  fontSize=24,
                                  textColor=RGBColor(0x26, 0x32, 0x44),
                                  spaceAfter=30,
                                  alignment=1))

        styles.add(ParagraphStyle(name='MetaLabel',
                                  parent=styles['Normal'],
                                  fontSize=10,
                                  textColor=RGBColor(0x64, 0x74, 0x8b),
                                  spaceAfter=6))

        styles.add(ParagraphStyle(name='MetaValue',
                                  parent=styles['Normal'],
                                  fontSize=12,
                                  spaceAfter=12))

        # Title
        title = Paragraph(script_data.get("title", "YouTube Script"), styles['CustomTitle'])
        elements.append(title)
        elements.append(Spacer(1, 12))

        # Metadata
        metadata_items = [
            ("Topic", script_data.get("topic", "N/A")),
            ("Style", script_data.get("tone", "N/A")),
            ("Duration", script_data.get("estimated_duration", "N/A"))
        ]

        for label, value in metadata_items:
            elements.append(Paragraph(f"<b>{label}</b>", styles['MetaLabel']))
            elements.append(Paragraph(value, styles['MetaValue']))

        elements.append(Spacer(1, 20))

        # Description
        if script_data.get("description"):
            elements.append(Paragraph("<b>Description</b>", styles['Heading2']))
            elements.append(Spacer(1, 6))
            elements.append(Paragraph(script_data.get("description"), styles['Normal']))
            elements.append(Spacer(1, 12))

        # Keywords
        if script_data.get("keywords"):
            elements.append(Paragraph("<b>Keywords</b>", styles['Heading2']))
            elements.append(Spacer(1, 6))
            elements.append(Paragraph(", ".join(script_data.get("keywords", [])), styles['Normal']))
            elements.append(Spacer(1, 12))

        # Script
        elements.append(Paragraph("<b>Script</b>", styles['Heading2']))
        elements.append(Spacer(1, 12))

        script_text = script_data.get("full_script", "")
        for paragraph in script_text.split("\n\n"):
            if paragraph.strip():
                elements.append(Paragraph(paragraph.strip(), styles['Normal']))
                elements.append(Spacer(1, 12))

        # Build PDF
        doc.build(elements)

        buffer.seek(0)
        return buffer.read()


export_service = ExportService()
